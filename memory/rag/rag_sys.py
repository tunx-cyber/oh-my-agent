# from transformers import AutoTokenizer, AutoModel
# import torch
# # 从transformers上下载模型
# class Embedding:
#     def __init__(
#         self,
#         hf_name: str = "BAAI/bge-small-zh"
#     ):
#         self.model_name = hf_name
#         self.model = None
#         self.tokenizer = None
    
#     def _load(self):
#         self.model = AutoModel.from_pretrained(self.model_name)
#         self.tokenizer = AutoTokenizer.from_pretrained(self.model_name)
    
#     def encode(self, texts):
#         if isinstance(texts, str):
#             texts = [texts]
#         tokenized = self.tokenizer(texts, return_tensors="pt", padding = True, trunction=True, max_length=512)
#         with torch.no_grad():
#             outputs = self.model(**tokenized)
#             embeddings = outputs.last_hidden_state.mean(dim=1).cpu().numpy()
        
#         vecs = [v for v in embeddings]
#         return vecs

# class Documentchunk:
#     def __init__(self):
#         pass

# class DocumentProcessor:
#     def __init__(self):
#         pass

#     def load_str(self,text):
#         pass

#     def load_file(self, file_path):
#         pass

#     def get_chunk(self):
#         pass

# class RagSystem:
#     def __init__(self):
#         pass

#     def add_document(self):
#         pass

#     def search(self):
#         pass

import os
import sys
from transformers import AutoTokenizer, AutoModel,AutoModelForCausalLM
import torch
from typing import List
import numpy as np
from pathlib import Path
import faiss
import csv
from typing import List, Dict, Any, Optional, Union
__package__ = "episode_2_RAG"
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

class Embedder:
    def __init__(self, model_name='BAAI/bge-small-zh'):
        self.tokenizer = AutoTokenizer.from_pretrained(model_name)
        self.model = AutoModel.from_pretrained(model_name)
        if torch.cuda.is_available():
            device = "cuda:0"
        elif torch.backends.mps.is_available():
            device = "mps"
        else:
            device = "cpu"
        self.device = device
        self.model.to(self.device)
        self.model.eval()
    
    def embed(self, texts: List[str]) -> np.ndarray:
        """将文本列表转换为嵌入向量"""
        embeddings = []
        
        for text in texts:
            # 标记化
            inputs = self.tokenizer(
                text, 
                return_tensors='pt', 
                truncation=True, 
                padding=True, 
                max_length=512
            ).to(self.device)
            
            # 前向传播
            with torch.no_grad():
                outputs = self.model(**inputs)
            
            # 使用平均池化获取句子嵌入
            embedding = outputs.last_hidden_state.mean(dim=1).squeeze().cpu().numpy()
            embeddings.append(embedding)
        
        return np.array(embeddings).astype('float32')

class UniversalFileReader:
    """通用文件读取器，支持多种格式"""
    
    def __init__(self, encoding: str = 'utf-8'):
        self.encoding = encoding
        self.supported_extensions = {
            '.txt', '.md', '.markdown',
            '.pdf', 
            '.doc', '.docx',
            '.csv'
        }
    
    def read(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        读取文件并返回结构化内容
        
        Returns:
            Dict with keys: 'content', 'metadata', 'pages' (for multi-page docs)
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 获取文件信息
        ext = file_path.suffix.lower()
        
        if ext == '.txt':
            return self._read_txt(file_path)
        elif ext in ['.md', '.markdown']:
            return self._read_markdown(file_path)
        elif ext == '.pdf':
            return self._read_pdf(file_path)
        elif ext in ['.doc', '.docx']:
            return self._read_docx(file_path)
        elif ext == '.csv':
            return self._read_csv(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
    
    def _read_txt(self, file_path: Path) -> Dict[str, Any]:
        """读取纯文本文件"""
        with open(file_path, 'r', encoding=self.encoding, errors='ignore') as f:
            content = f.read()
        
        return {
            'content': content,
            'metadata': {
                'file_type': 'txt',
                'file_size': os.path.getsize(file_path)
            },
            'pages': [{'page_num': 1, 'content': content}]
        }
    
    def _read_markdown(self, file_path: Path) -> Dict[str, Any]:
        """读取Markdown文件"""
        with open(file_path, 'r', encoding=self.encoding) as f:
            md_content = f.read()
        import markdown
        # 可选：将Markdown转换为纯文本（保留结构信息）
        html_content = markdown.markdown(md_content)
        
        return {
            'content': md_content,  # 原始Markdown
            'clean_content': self._clean_markdown(md_content),  # 清理后的文本
            'metadata': {
                'file_type': 'markdown',
                'has_headers': '#' in md_content,
            },
            'pages': [{'page_num': 1, 'content': md_content}]
        }
    
    def _read_pdf(self, file_path: Path) -> Dict[str, Any]:
        """读取PDF文件（使用pdfplumber提取文本和表格）"""
        pages_content = []
        full_text = ""
        metadata = {}
        import pdfplumber
        try:
            # 方法1：使用pdfplumber（更好的文本提取）
            with pdfplumber.open(file_path) as pdf:
                metadata = {
                    'file_type': 'pdf',
                    'total_pages': len(pdf.pages),
                    'author': pdf.metadata.get('Author', ''),
                    'title': pdf.metadata.get('Title', '')
                }
                
                for i, page in enumerate(pdf.pages):
                    # 提取文本
                    text = page.extract_text() or ""
                    # 提取表格
                    tables = page.extract_tables()
                    
                    pages_content.append({
                        'page_num': i + 1,
                        'content': text,
                        'tables': tables,
                        'bbox': page.bbox
                    })
                    full_text += text + "\n\n"
                
                # 方法2：备用方案使用PyPDF2
                if not full_text.strip():
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        full_text = ""
                        for i, page in enumerate(pdf_reader.pages):
                            text = page.extract_text() or ""
                            full_text += text + "\n\n"
        
        except Exception as e:
            raise Exception(f"PDF读取失败: {str(e)}")
        
        return {
            'content': full_text,
            'metadata': metadata,
            'pages': pages_content,
        }
    
    def _read_docx(self, file_path: Path) -> Dict[str, Any]:
        """读取Word文档"""
        from docx import Document
        try:
            doc = Document(file_path)
            full_text = []
            metadata = {
                'file_type': 'docx',
                'paragraphs_count': len(doc.paragraphs)
            }
            
            # 提取段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # 提取表格
            tables_content = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    tables_content.append(" | ".join(row_text))
            
            content = "\n".join(full_text)
            if tables_content:
                content += "\n\n表格:\n" + "\n".join(tables_content)
            
            return {
                'content': content,
                'metadata': metadata,
                'pages': [{'page_num': 1, 'content': content}]
            }
        except Exception as e:
            raise Exception(f"Word文档读取失败: {str(e)}")
    
    def _read_csv(self, file_path: Path) -> Dict[str, Any]:
        """读取CSV文件"""
        import pandas as pd
        try:
            # 尝试自动检测编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            content = None
            
            for enc in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=enc)
                    content = df.to_string()
                    break
                except:
                    continue
            
            if content is None:
                # 如果自动检测失败，使用简单文本读取
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            
            return {
                'content': content,
                'metadata': {
                    'file_type': 'csv',
                    'delimiter': ',',
                    'rows': len(pd.read_csv(file_path)) if 'df' in locals() else 'unknown'
                },
                'pages': [{'page_num': 1, 'content': content}],
                'language': 'en'  # CSV通常是数据，语言检测意义不大
            }
        except Exception as e:
            raise Exception(f"CSV读取失败: {str(e)}")
    
    def _clean_markdown(self, text: str) -> str:
        """清理Markdown格式，保留主要内容"""
        import re
        # 移除代码块
        text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)
        # 移除行内代码
        text = re.sub(r'`.*?`', '', text)
        # 移除图片链接
        text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
        # 移除普通链接但保留文字
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
        # 移除标题标记
        text = re.sub(r'#+ ', '', text)
        return text.strip()
    

class SemanticChunker:
    """基于语义的智能文档分块器"""
    
    def __init__(self, 
                 chunk_size: int = 1000,
                 chunk_overlap: int = 200):
        """
        初始化分块器
        
        Args:
            chunk_size: 每个块的目标大小（字符数）
            chunk_overlap: 块之间的重叠大小
            model_name: 用于token计数的模型名称
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # 定义分割符优先级（按语义边界）
        self.separators = [
            "\n\n",  # 段落边界（最高优先级）
            "\n",    # 换行
            ". ",    # 句子边界
            "! ", "? ",  # 其他句子结束符
            "; ", ": ", ", ",  # 子句边界
            " ",     # 单词边界（最后选择）
            ""       # 字符边界（兜底）
        ]
    
    def recursive_split(self, text: str, separators: list = None) -> List[str]:
        """递归字符分割（LangChain风格的最佳实践）"""
        if separators is None:
            separators = self.separators
        
        # 最终分割结果
        chunks = []
        
        # 选择合适的分隔符
        separator = self._choose_separator(text, separators)
        
        if separator:
            splits = text.split(separator)
            
            # 合并小片段
            current_chunk = ""
            for split in splits:
                # 如果当前块加上新分割部分太大，则保存当前块
                if len(current_chunk) + len(split) + len(separator) > self.chunk_size:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    
                    # 如果单个分割部分就超过chunk_size，递归分割它
                    if len(split) > self.chunk_size:
                        sub_chunks = self.recursive_split(
                            split, 
                            separators[separators.index(separator) + 1:] if separator in separators else []
                        )
                        chunks.extend(sub_chunks[:-1])  # 添加除最后一个外的所有子块
                        current_chunk = sub_chunks[-1] if sub_chunks else ""
                    else:
                        current_chunk = split + separator
                else:
                    current_chunk += split + separator
            
            # 添加最后一个块
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
        else:
            # 如果没有合适的分隔符，按长度硬分割
            if len(text) <= self.chunk_size:
                chunks.append(text)
            else:
                chunks.append(text[:self.chunk_size])
                chunks.extend(self.recursive_split(
                    text[self.chunk_size - self.chunk_overlap:],
                    separators
                ))
        
        # 应用重叠
        if self.chunk_overlap > 0 and len(chunks) > 1:
            chunks = self._add_overlap(chunks)
        
        return chunks
    
    def _choose_separator(self, text: str, separators: list) -> Optional[str]:
        """选择最合适的分隔符"""
        for separator in separators:
            if separator in text:
                return separator
        return None
    
    def _add_overlap(self, chunks: List[str]) -> List[str]:
        """在块之间添加重叠内容"""
        overlapped_chunks = []
        
        for i, chunk in enumerate(chunks):
            if i > 0:
                # 从前一个块末尾取重叠部分
                prev_chunk = chunks[i-1]
                overlap_start = max(0, len(prev_chunk) - self.chunk_overlap)
                overlap_text = prev_chunk[overlap_start:]
                
                # 将重叠部分添加到当前块开头
                chunk = overlap_text + "\n" + chunk
            
            overlapped_chunks.append(chunk)
        
        return overlapped_chunks
    
    def split_document(self, document: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        分割文档，保留元数据
        
        Args:
            document: UniversalFileReader返回的文档字典
            
        Returns:
            List of chunks with metadata
        """
        content = document.get('content', '')
        metadata = document.get('metadata', {})
        pages = document.get('pages', [])
        
        # 根据不同文件类型采用不同策略
        file_type = metadata.get('file_type', 'unknown')
        
        if file_type == 'pdf' and pages:
            # PDF按页面分割，再在页面内分块
            all_chunks = []
            for page in pages:
                page_content = page.get('content', '')
                page_num = page.get('page_num', 1)
                
                page_chunks = self.recursive_split(page_content)
                
                for i, chunk in enumerate(page_chunks):
                    chunk_metadata = metadata.copy()
                    chunk_metadata.update({
                        'chunk_id': f"page_{page_num}_chunk_{i+1}",
                        'page_number': page_num,
                        'chunk_index': i,
                        'total_chunks_in_page': len(page_chunks),
                        'source': f"PDF Page {page_num}"
                    })
                    all_chunks.append({
                        'content': chunk,
                        'metadata': chunk_metadata
                    })
            
            return all_chunks
        
        elif file_type == 'markdown':
            # Markdown按标题分割
            return self._split_by_markdown_headers(content, metadata)
        
        else:
            # 通用递归分割
            text_chunks = self.recursive_split(content)
            
            chunks_with_meta = []
            for i, chunk in enumerate(text_chunks):
                chunk_metadata = metadata.copy()
                chunk_metadata.update({
                    'chunk_id': f"chunk_{i+1}",
                    'chunk_index': i,
                    'total_chunks': len(text_chunks)
                })
                
                chunks_with_meta.append({
                    'content': chunk,
                    'metadata': chunk_metadata
                })
            
            return chunks_with_meta
    
    def _split_by_markdown_headers(self, text: str, metadata: Dict) -> List[Dict[str, Any]]:
        """按Markdown标题分割（保持文档结构）"""
        import re
        
        # 匹配Markdown标题（# Header, ## Subheader等）
        header_pattern = r'(^#+\s+.+$)'
        
        splits = re.split(header_pattern, text, flags=re.MULTILINE)
        
        chunks = []
        current_header = "文档开头"
        current_content = []
        
        for i, part in enumerate(splits):
            if re.match(header_pattern, part.strip()):
                # 保存前一个块
                if current_content:
                    chunk_text = current_header + "\n\n" + "\n".join(current_content)
                    text_chunks = self.recursive_split(chunk_text)
                    
                    for j, sub_chunk in enumerate(text_chunks):
                        chunk_metadata = metadata.copy()
                        chunk_metadata.update({
                            'chunk_id': f"section_{len(chunks)+1}_sub_{j+1}",
                            'section_header': current_header,
                            'chunk_index': j
                        })
                        chunks.append({
                            'content': sub_chunk,
                            'metadata': chunk_metadata
                        })
                
                # 开始新块
                current_header = part.strip()
                current_content = []
            else:
                if part.strip():
                    current_content.append(part.strip())
        
        # 处理最后一个块
        if current_content:
            chunk_text = current_header + "\n\n" + "\n".join(current_content)
            text_chunks = self.recursive_split(chunk_text)
            
            for j, sub_chunk in enumerate(text_chunks):
                chunk_metadata = metadata.copy()
                chunk_metadata.update({
                    'chunk_id': f"section_{len(chunks)+1}_sub_{j+1}",
                    'section_header': current_header,
                    'chunk_index': j
                })
                chunks.append({
                    'content': sub_chunk,
                    'metadata': chunk_metadata
                })
        
        return chunks
        
class RAGSys:
    def __init__(self, data_dir, embedder:Embedder):
        self.documents = self.process_data(data_dir)
        self.embedder = embedder
        self.embeddings = embedder.embed(self.documents)
        self.dimension = self.embeddings[0].shape[-1]
        print(f"Dimension for faiss is {self.dimension}")
        self.index = faiss.IndexHNSWFlat(self.dimension, 32)
        print("创建HNSW索引")
        # 设置构建参数
        self.index.hnsw.efConstruction = 20  # 构建时考虑的邻居数量
        self.index.hnsw.efSearch = 50         # 搜索时考虑的邻居数量    
        # 添加数据
        self.index.add(self.embeddings)

    def get_all_files_pathlib(self, directory):
        """使用 pathlib.Path.rglob 获取目录下所有文件的路径"""
        dir_path = Path(directory)
        if dir_path.is_absolute():
            path = dir_path
        else:
            base_dir = Path(__file__).resolve().parent
            path = Path(base_dir/directory)
        # 使用 ‘*’ 匹配所有文件，递归地
        all_files = list(path.rglob("*"))
        # 通常我们只想要文件，可以过滤掉目录
        all_files = [p for p in all_files if p.is_file()]
        return all_files
    
    def process_data(self, dir):
        file_list = self.get_all_files_pathlib(dir)
        file_contents = []
        file_reader = UniversalFileReader()
        for file in file_list:
            file_contents.append(file_reader.read(file.as_posix()))
        chunker = SemanticChunker(chunk_size=200, chunk_overlap=150)
        file_chunks = []
        for file_content in file_contents:
            chunks = chunker.split_document(file_content)
            for i, chunk in enumerate(chunks):
                file_chunks.append(chunk['content'])
        return file_chunks

    def search(self, query):
        query = self.embedder.embed(query)
        distances, indices = self.index.search(query, 5)
        print(indices)
        return [self.documents[i] for i in indices[0]]


# prompt = f"根据以下信息回答问题:\n{context}\n\n问题: {question}\n"
embedder = Embedder()
rag = RAGSys("./data",embedder)
res = rag.search(["请问哪一项实验室检查对了解胸水的性质更有帮助？"])
print(res)
model_name = "Qwen/Qwen3-0.6B"
tokenizer = AutoTokenizer.from_pretrained(model_name)
if torch.cuda.is_available():
    device = "cuda:0"
elif torch.backends.mps.is_available():
    device = "mps"
else:
    device = "cpu"
model = AutoModelForCausalLM.from_pretrained(model_name,dtype="auto").to(device)
messages = [
    {"role": "system", "content": "You are a help AI assitant"},
    {"role": "user", "content": f"根据以下信息回答问题:\n{res}\n\n问题:请问哪一项实验室检查对了解胸水的性质更有帮助？"}
]
prompt = tokenizer.apply_chat_template(
                messages,
                tokenize=False,
                add_generation_prompt=True,
                enable_thinking=True # Switches between thinking and non-thinking modes. Default is True.
            )
enc = tokenizer(prompt,
        truncation=True,
        return_tensors='pt').to(device)
from transformers import TextIteratorStreamer
streamer = TextIteratorStreamer(tokenizer, skip_prompt=True, timeout=60.0, skip_special_tokens=True)
generation_kwargs = dict(
    **enc,
    streamer=streamer,
    max_new_tokens=32768,
    do_sample=True,
    temperature=0.7,
    pad_token_id=tokenizer.eos_token_id,
    eos_token_id=tokenizer.eos_token_id
)
from threading import Thread
thread = Thread(target=model.generate, kwargs=generation_kwargs)
thread.start()
for new_text in streamer:
    # full_response += new_text
    print(new_text,end='',flush=True)